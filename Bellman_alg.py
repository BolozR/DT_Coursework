import matplotlib.pyplot as plt
import math
from prettytable import PrettyTable
import docx


class Bellman_Alg(object):

    def __init__(self, g_input, x_input, step_num, previous=[0], disp=True):
        self.g_input = g_input
        self.x_input = x_input
        self.step_num = step_num
        self.previous = previous
        self.disp = disp

    def plot(self, gArr, x4Arr, matrix, labels, max_columns, titles):
        plt.title(titles[0])
        plt.xlabel(labels[0])
        plt.ylabel(labels[1])
        plt.grid()
        plt.plot(x4Arr, matrix)
        plt.legend(gArr)
        plt.show()
        plt.title(titles[1])
        plt.xlabel(labels[2])
        plt.ylabel(labels[3])
        plt.grid()
        plt.plot(gArr, max_columns)
        plt.show()

# для поиска максимума целевой функции:
    def max_in_matrix_column(self, matrix):
        max = [-math.inf for _ in range(len(matrix[0]))]
        for i in range(len(matrix[0])):
            for j in range(len(matrix)):
                if matrix[j][i] >= max[i]:
                    max[i] = matrix[j][i]
                    for k in max:
                        if max[i] >= k:
                            self.max_i = j
        print("Max elements:\n", max)
        return max

# расчёт целевой функции:
    def func_z(self, gArr, xArr, param):
        z = lambda i, j: 200 + 1.4 * xArr[j] ** (3 / 4) + 0.7 * (gArr[j] - xArr[i]) ** (3 / 4)
        matrix = [[0 for _ in range(len(gArr))] for _ in range(len(xArr))]
        for i in range(len(matrix)):
            for j in range(len(matrix[0])):
                count = z(i, j)
                if isinstance(count, complex) and (count.imag != 0):
                    count = float('nan')
                matrix[i][j] = round(count + param[i][j], 2)
        return matrix

# таблица со значениями целевой функции в doc файл:
    def tabel_in_doc(self, field, x, matr, last=False):
        doc = docx.Document()
        if not last: table_doc = doc.add_table(rows=12, cols=7)
        else: table_doc = doc.add_table(rows=12, cols=2)
        table_doc.style = 'Table Grid'
        for i in range(len(field)):
            cell = table_doc.cell(0, i)
            cell.text = field[i]
        for i in range(len(x)):
            cell = table_doc.cell(i+1, 0)
            cell.text = str(round(x[i], 2))
        for i in range(len(matr)):
            if not last:
                for j in range(len(matr[0])):
                    cell = table_doc.cell(i + 1, j + 1)
                    cell.text = str(round(matr[i][j], 2))
            else:
                cell = table_doc.cell(i + 1, 1)
                cell.text = str(round(matr[i], 2))
        doc.save("table_task" + str(self.step_num) + ".docx")

# шаг алгоритма:
    def step(self):
        xArr = []
        gArr = []
        # массив x c шагом x_input:
        for i in range(11):
            xArr.append(round(self.x_input * i, 2))
        # массив g c шагом g_input
        for i in range(6):
            gArr.append(round(self.g_input * i + xArr[-1], 2))
        table = PrettyTable()
        fields = ["x\\g"]
        for i in gArr:
            fields.append(str(i))
        table.field_names = fields
        zeroArr = [[0 for _ in range(len(gArr))] for _ in range(len(xArr))]
        param = zeroArr
        new_gArr = [0 for _ in range(len(gArr))]
        # расчёт целевой функции предыдущего шага:
        if self.step_num != 4:
            for i in range(len(zeroArr[0])):
                new_gArr[i] = (round((0.64 * gArr[i]), 2))
            param = self.func_z(new_gArr, xArr, zeroArr)
        # расчёт целевой функции:
        matrix = self.func_z(gArr, xArr, param)
        for i in range(len(xArr)):
            table.add_row([xArr[i]] + matrix[i])
        print(table)
        if(self.disp): self.tabel_in_doc(fields, xArr, matrix)
        self.max = self.max_in_matrix_column(matrix)
        self.max_x = xArr[self.max_i]
        if self.disp:
            self.plot(gArr, xArr, matrix,
                      ["x" + str(self.step_num), "W" + str(self.step_num), "g", "W*" + str(self.step_num) + "(g)"],
                      self.max, ["Step " + str(self.step_num), "Step " + str(self.step_num)])
        return matrix

# последний шаг боратного алгоритма:
    def last_step(self, g):
        xArr = []
        for i in range(11):
            xArr.append(round(self.x_input * i, 2))
        table = PrettyTable()
        table.field_names = ["x\\g", str(g)]
        zeroArr = [0 for _ in range(len(xArr))]
        param = zeroArr
        new_g = (round((0.19 * 0 + 0.64 * g), 2))

        z = lambda i, j: 200 + 1.4 * xArr[i] ** (3 / 4) + 0.7 * (j - xArr[i]) ** (3 / 4)
        for i in range(len(param)):
            count = z(i, new_g)
            if isinstance(count, complex) and (count.imag != 0):
                count = float('nan')
            param[i] = round(count, 2)
        matrix = [0 for _ in range(len(xArr))]
        for i in range(len(matrix)):
            count = z(i, g)
            if isinstance(count, complex) and (count.imag != 0):
                count = float('nan')
            matrix[i] = round(count + param[i], 2)
            table.add_row([str(xArr[i]), str(matrix[i])])

        print(table)
        if (self.disp): self.tabel_in_doc(["x\\g", str(g)], xArr, matrix, True)

        self.max = -math.inf
        self.max_i = -math.inf
        for i in range(len(matrix)):
            if self.max <= matrix[i]:
                self.max = matrix[i]
                self.max_i = i
        print("max: " + str(self.max))
        self.max_x = xArr[self.max_i]
        return matrix
